"""One-off generator for Phase 2 smoke-test fixtures.

Run from repo root inside a venv with:
    pip install reportlab python-docx
    python backend/scripts/generate_smoke_fixtures.py

Writes backend/tests/fixtures/hello.pdf and hello.docx (<50KB each).
These fixtures are consumed by backend/scripts/docker_smoke.sh to prove
PDF + DOCX ingest inside the container works (DEPLOY-01 criterion 3).
NOTE: reportlab and python-docx are NOT shipped in the image -- this is a
host-side regeneration utility only.
"""
from pathlib import Path

from reportlab.pdfgen import canvas
from docx import Document

FIXTURES = Path(__file__).resolve().parent.parent / "tests" / "fixtures"
FIXTURES.mkdir(parents=True, exist_ok=True)

pdf_path = FIXTURES / "hello.pdf"
c = canvas.Canvas(str(pdf_path))
c.drawString(100, 750, "Hello from Phase 2 smoke-test fixture.")
c.drawString(100, 730, "This PDF verifies Docling PDF ingest inside the container.")
c.showPage()
c.save()
print(f"Wrote {pdf_path} ({pdf_path.stat().st_size} bytes)")

docx_path = FIXTURES / "hello.docx"
d = Document()
d.add_paragraph("Hello from Phase 2 smoke-test fixture.")
d.add_paragraph("This DOCX verifies Docling DOCX ingest inside the container.")
d.save(str(docx_path))
print(f"Wrote {docx_path} ({docx_path.stat().st_size} bytes)")
